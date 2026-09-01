"""Export functionality for Cognitia Brain."""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


class Exporter:
    """Export fichamentos in multiple formats."""

    def __init__(self, resumos_dir: Path = Path("resumos")):
        self.resumos_dir = resumos_dir

    def export_markdown(self, filename: str, output_path: Optional[Path] = None) -> Path:
        """Export fichamento as Markdown."""
        source_path = self.resumos_dir / filename
        if not source_path.exists():
            raise FileNotFoundError(f"Fichamento não encontrado: {filename}")

        content = source_path.read_text(encoding="utf-8")

        if output_path is None:
            output_path = Path(f"exports/{filename}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")

        return output_path

    def export_pdf(self, filename: str, output_path: Optional[Path] = None) -> Path:
        """Export fichamento as PDF."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
        except ImportError:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")

        source_path = self.resumos_dir / filename
        if not source_path.exists():
            raise FileNotFoundError(f"Fichamento não encontrado: {filename}")

        content = source_path.read_text(encoding="utf-8")

        if output_path is None:
            output_path = Path(f"exports/{filename.replace('.md', '.pdf')}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()

        # Custom style for title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12
        )

        # Build content
        story = []
        lines = content.split('\n')

        for line in lines:
            if line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.strip():
                story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.1 * inch))

        doc.build(story)

        return output_path

    def export_docx(self, filename: str, output_path: Optional[Path] = None) -> Path:
        """Export fichamento as DOCX."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

        source_path = self.resumos_dir / filename
        if not source_path.exists():
            raise FileNotFoundError(f"Fichamento não encontrado: {filename}")

        content = source_path.read_text(encoding="utf-8")

        if output_path is None:
            output_path = Path(f"exports/{filename.replace('.md', '.docx')}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create DOCX
        doc = Document()

        # Add title
        title = doc.add_heading(filename.replace('.md', '').replace('_', ' '), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse content
        lines = content.split('\n')
        current_section = None

        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(str(output_path))

        return output_path

    def export_bibtex(self, filename: str, output_path: Optional[Path] = None) -> Path:
        """Export fichamento as BibTeX."""
        source_path = self.resumos_dir / filename
        if not source_path.exists():
            raise FileNotFoundError(f"Fichamento não encontrado: {filename}")

        content = source_path.read_text(encoding="utf-8")

        if output_path is None:
            output_path = Path(f"exports/{filename.replace('.md', '.bib')}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Extract metadata from content
        title = filename.replace('.md', '').replace('_', ' ')
        author = "Cognitia Brain"
        year = datetime.now().year

        # Simple BibTeX entry
        bibtex = f"""@article{{{filename.replace('.md', '').replace(' ', '_')},
  title = {{{title}}},
  author = {{{author}}},
  year = {{{year}}},
  note = {{Fichamento gerado automaticamente pelo Cognitia Brain}}
}}
"""

        output_path.write_text(bibtex, encoding="utf-8")

        return output_path

    def export_json(self, filename: str, output_path: Optional[Path] = None) -> Path:
        """Export fichamento as JSON."""
        source_path = self.resumos_dir / filename
        if not source_path.exists():
            raise FileNotFoundError(f"Fichamento não encontrado: {filename}")

        content = source_path.read_text(encoding="utf-8")

        if output_path is None:
            output_path = Path(f"exports/{filename.replace('.md', '.json')}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Parse content into structured data
        data = {
            "filename": filename,
            "content": content,
            "exported_at": datetime.now().isoformat(),
            "format": "json"
        }

        output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

        return output_path

    def export_all_formats(self, filename: str, output_dir: Path = Path("exports")) -> Dict[str, Path]:
        """Export fichamento in all available formats."""
        output_dir.mkdir(parents=True, exist_ok=True)

        results = {}

        # Markdown
        try:
            results["markdown"] = self.export_markdown(filename, output_dir / filename)
        except Exception as e:
            logger.error(f"Error exporting markdown: {e}")

        # PDF
        try:
            results["pdf"] = self.export_pdf(filename, output_dir / filename.replace('.md', '.pdf'))
        except Exception as e:
            logger.error(f"Error exporting PDF: {e}")

        # DOCX
        try:
            results["docx"] = self.export_docx(filename, output_dir / filename.replace('.md', '.docx'))
        except Exception as e:
            logger.error(f"Error exporting DOCX: {e}")

        # BibTeX
        try:
            results["bibtex"] = self.export_bibtex(filename, output_dir / filename.replace('.md', '.bib'))
        except Exception as e:
            logger.error(f"Error exporting BibTeX: {e}")

        # JSON
        try:
            results["json"] = self.export_json(filename, output_dir / filename.replace('.md', '.json'))
        except Exception as e:
            logger.error(f"Error exporting JSON: {e}")

        return results


# Global exporter instance
exporter = Exporter()
